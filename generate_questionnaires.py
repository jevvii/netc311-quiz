import os
import random
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from module1_data import MODULE_1_ITEMS
from module2_data import MODULE_2_ITEMS

# Styling Palette (Professional Cisco Academic Theme)
COLOR_PRIMARY = RGBColor(0, 51, 102)     # Deep Cisco Navy
COLOR_SECONDARY = RGBColor(0, 102, 153)  # Tech Blue
COLOR_DARK = RGBColor(34, 34, 34)        # Charcoal text
COLOR_GRAY = RGBColor(100, 100, 100)     # Subtle gray
HEX_PRIMARY = "003366"
HEX_LIGHT_BG = "F4F6F9"
HEX_BORDER = "D0D7DE"

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D0D7DE"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def build_module_doc(module_title, module_subtitle, module_code, items, output_docx_path):
    doc = docx.Document()

    # Configure Margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Document Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(2)
    run_course = header_para.add_run("NETC311: INTRODUCTION TO NETWORKS v7.0 (ITN)")
    run_course.font.name = 'Arial'
    run_course.font.size = Pt(11)
    run_course.font.bold = True
    run_course.font.color.rgb = COLOR_SECONDARY

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    run_title = title_para.add_run(module_title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(12)
    run_sub = sub_para.add_run(module_subtitle + " — Identification Questionnaire")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_GRAY

    # Student Info Box / Table
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    col_widths = [Inches(4.25), Inches(2.75)]
    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_background(cell, "F8F9FA")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    # Info text
    p = info_table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Name: _________________________________________")
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_DARK

    p = info_table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Date: ________________________")
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_DARK

    p = info_table.cell(1, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Course & Year: NETC311 / Year III")
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_DARK

    p = info_table.cell(1, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Score: _________ / {len(items)}")
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = COLOR_PRIMARY

    set_table_borders(info_table, color="D0D7DE")

    # Instructions Section
    inst_para = doc.add_paragraph()
    inst_para.paragraph_format.space_before = Pt(12)
    inst_para.paragraph_format.space_after = Pt(14)
    run_inst_title = inst_para.add_run("GENERAL INSTRUCTIONS: ")
    run_inst_title.font.name = 'Arial'
    run_inst_title.font.size = Pt(9.5)
    run_inst_title.font.bold = True
    run_inst_title.font.color.rgb = COLOR_PRIMARY

    run_inst = inst_para.add_run(
        f"This questionnaire contains {len(items)} identification-type items based thoroughly on {module_title}. "
        "Read each statement carefully. Identify the correct term that completes the blank (____) by selecting the best answer "
        "among choices a, b, c, or d. Write the letter of your choice on the blank provided before each number or encircle the letter. "
        "A complete Answer Key and Study Reference is provided at the end of this document."
    )
    run_inst.font.name = 'Arial'
    run_inst.font.size = Pt(9.5)
    run_inst.font.color.rgb = COLOR_DARK

    # Divider Heading: PART I
    part1_para = doc.add_paragraph()
    part1_para.paragraph_format.space_before = Pt(8)
    part1_para.paragraph_format.space_after = Pt(8)
    run_part1 = part1_para.add_run("PART I: IDENTIFICATION QUESTIONNAIRE")
    run_part1.font.name = 'Arial'
    run_part1.font.size = Pt(12)
    run_part1.font.bold = True
    run_part1.font.color.rgb = COLOR_PRIMARY

    # Process items deterministically
    processed_items = []
    rng = random.Random(101 + int(module_code))

    for idx, item in enumerate(items, 1):
        q_text = item["q"]
        correct_answer = item["a"]
        distractors = item["distractors"]
        
        # 4 options: correct answer + 3 distractors
        all_options = [correct_answer] + distractors
        # Deterministic shuffle for this question
        rng.shuffle(all_options)
        
        # Find index of correct answer
        correct_idx = all_options.index(correct_answer)
        letters = ['a', 'b', 'c', 'd']
        correct_letter = letters[correct_idx]

        processed_items.append({
            "number": idx,
            "q": q_text,
            "options": all_options,
            "correct_letter": correct_letter,
            "correct_answer": correct_answer,
            "topic": item["topic"],
            "explanation": item["explanation"]
        })

    # Render questions in DOCX
    for pitem in processed_items:
        idx = pitem["number"]
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(6)
        q_para.paragraph_format.space_after = Pt(2)
        q_para.paragraph_format.left_indent = Inches(0.0)

        # Question prefix: "____ 1. Question text"
        run_blank = q_para.add_run("_____ ")
        run_blank.font.name = 'Arial'
        run_blank.font.size = Pt(10)
        run_blank.font.bold = True
        run_blank.font.color.rgb = COLOR_PRIMARY

        run_num = q_para.add_run(f"{idx}. ")
        run_num.font.name = 'Arial'
        run_num.font.size = Pt(10)
        run_num.font.bold = True
        run_num.font.color.rgb = COLOR_DARK

        run_q = q_para.add_run(pitem["q"])
        run_q.font.name = 'Arial'
        run_q.font.size = Pt(10)
        run_q.font.color.rgb = COLOR_DARK

        # Render options in 2x2 grid or clean indented lines
        # Using 2 lines with 2 options each or 4 indented lines
        # Check maximum option length
        max_opt_len = max(len(opt) for opt in pitem["options"])
        
        letters = ['a', 'b', 'c', 'd']
        if max_opt_len < 28:
            # 2 lines, 2 columns format
            opt_p1 = doc.add_paragraph()
            opt_p1.paragraph_format.left_indent = Inches(0.35)
            opt_p1.paragraph_format.space_before = Pt(1)
            opt_p1.paragraph_format.space_after = Pt(1)

            r = opt_p1.add_run(f"a.) {pitem['options'][0]}")
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = COLOR_DARK
            
            # spacing tab
            r_space = opt_p1.add_run(" " * (35 - len(pitem['options'][0])))
            r_space.font.name = 'Arial'
            r_space.font.size = Pt(9.5)

            r = opt_p1.add_run(f"b.) {pitem['options'][1]}")
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = COLOR_DARK

            opt_p2 = doc.add_paragraph()
            opt_p2.paragraph_format.left_indent = Inches(0.35)
            opt_p2.paragraph_format.space_before = Pt(1)
            opt_p2.paragraph_format.space_after = Pt(4)

            r = opt_p2.add_run(f"c.) {pitem['options'][2]}")
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = COLOR_DARK

            r_space = opt_p2.add_run(" " * (35 - len(pitem['options'][2])))
            r_space.font.name = 'Arial'
            r_space.font.size = Pt(9.5)

            r = opt_p2.add_run(f"d.) {pitem['options'][3]}")
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = COLOR_DARK
        else:
            # 4 lines format for longer options
            for opt_idx, letter in enumerate(letters):
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.left_indent = Inches(0.35)
                opt_p.paragraph_format.space_before = Pt(1)
                opt_p.paragraph_format.space_after = Pt(1 if opt_idx < 3 else 4)

                r_let = opt_p.add_run(f"{letter}.) ")
                r_let.font.name = 'Arial'
                r_let.font.size = Pt(9.5)
                r_let.font.bold = True
                r_let.font.color.rgb = COLOR_SECONDARY

                r_val = opt_p.add_run(pitem['options'][opt_idx])
                r_val.font.name = 'Arial'
                r_val.font.size = Pt(9.5)
                r_val.font.color.rgb = COLOR_DARK

    # Page Break for Answer Key Section
    doc.add_page_break()

    # Part II: Answer Key & Study Reference
    part2_para = doc.add_paragraph()
    part2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part2_para.paragraph_format.space_before = Pt(6)
    part2_para.paragraph_format.space_after = Pt(2)
    run_part2 = part2_para.add_run(f"PART II: ANSWER KEY & STUDY REFERENCE — {module_title.upper()}")
    run_part2.font.name = 'Arial'
    run_part2.font.size = Pt(14)
    run_part2.font.bold = True
    run_part2.font.color.rgb = COLOR_PRIMARY

    sub_key = doc.add_paragraph()
    sub_key.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_key.paragraph_format.space_after = Pt(14)
    r = sub_key.add_run("Official Solutions, Correct Terms, and Curriculum Concept References")
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COLOR_GRAY

    # Answer Key Table
    key_table = doc.add_table(rows=len(processed_items) + 1, cols=4)
    key_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    key_table.autofit = False

    table_widths = [Inches(0.6), Inches(0.7), Inches(2.2), Inches(3.5)]
    headers = ["Item #", "Key", "Correct Identification Term", "Curriculum Topic & Reference Explanation"]

    # Header Row
    hdr_cells = key_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = table_widths[i]
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for row_idx, pitem in enumerate(processed_items, 1):
        row_cells = key_table.rows[row_idx].cells
        bg_color = HEX_LIGHT_BG if row_idx % 2 == 0 else "FFFFFF"
        
        for i in range(4):
            row_cells[i].width = table_widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=60, bottom=60, left=80, right=80)

        # Col 0: Item #
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(pitem["number"]))
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = COLOR_DARK

        # Col 1: Correct Letter
        p = row_cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{pitem['correct_letter']}.)")
        r.font.name = 'Arial'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY

        # Col 2: Term
        p = row_cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(pitem["correct_answer"])
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = COLOR_DARK

        # Col 3: Explanation & Topic
        p = row_cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        
        r_top = p.add_run(f"[{pitem['topic']}] ")
        r_top.font.name = 'Arial'
        r_top.font.size = Pt(8.5)
        r_top.font.bold = True
        r_top.font.color.rgb = COLOR_SECONDARY

        r_exp = p.add_run(pitem["explanation"])
        r_exp.font.name = 'Arial'
        r_exp.font.size = Pt(8.5)
        r_exp.font.color.rgb = COLOR_DARK

    set_table_borders(key_table, color="D0D7DE")

    # Save document
    doc.save(output_docx_path)
    print(f"Generated DOCX: {output_docx_path} ({len(items)} items)")

print("Script template ready.")

if __name__ == "__main__":
    base_dir = "/home/javvii/YearIII/NETC311/quiz1"
    
    # Module 1
    m1_docx = os.path.join(base_dir, "Module 1 - Networking Today - Questionnaire.docx")
    build_module_doc(
        module_title="Module 1: Networking Today",
        module_subtitle="CCNA 1: Introduction to Networks v7.0",
        module_code="1",
        items=MODULE_1_ITEMS,
        output_docx_path=m1_docx
    )

    # Module 2
    m2_docx = os.path.join(base_dir, "Module 2 - Basic Switch and End Device Configuration - Questionnaire.docx")
    build_module_doc(
        module_title="Module 2: Basic Switch and End Device Configuration",
        module_subtitle="CCNA 1: Introduction to Networks v7.0",
        module_code="2",
        items=MODULE_2_ITEMS,
        output_docx_path=m2_docx
    )

    # Convert both to PDF via LibreOffice
    print("Converting documents to PDF via LibreOffice...")
    cmd1 = ["libreoffice", "--headless", "--convert-to", "pdf", m1_docx, "--outdir", base_dir]
    res1 = subprocess.run(cmd1, capture_output=True, text=True)
    print("Module 1 PDF conversion:", res1.returncode, res1.stdout.strip(), res1.stderr.strip())

    cmd2 = ["libreoffice", "--headless", "--convert-to", "pdf", m2_docx, "--outdir", base_dir]
    res2 = subprocess.run(cmd2, capture_output=True, text=True)
    print("Module 2 PDF conversion:", res2.returncode, res2.stdout.strip(), res2.stderr.strip())

    print("All documents generated successfully.")
