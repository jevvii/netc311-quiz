import os
import random
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from hmby311_m1_data import HMBY_MODULE_1_ITEMS

COLOR_PRIMARY = RGBColor(20, 82, 40)     # Biology Deep Forest Green
COLOR_SECONDARY = RGBColor(46, 125, 50)  # Leaf Green
COLOR_DARK = RGBColor(34, 34, 34)        # Charcoal text
COLOR_GRAY = RGBColor(100, 100, 100)     # Subtle gray
HEX_PRIMARY = "145228"
HEX_LIGHT_BG = "F4F8F4"
HEX_BORDER = "C8D6C8"

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="C8D6C8"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def build_hmby_doc(module_title, module_subtitle, module_code, items, output_docx_path):
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Document Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(2)
    run_course = header_para.add_run("HMBY311: HUMAN BIOLOGY")
    run_course.font.name = 'Arial'
    run_course.font.size = Pt(11)
    run_course.font.bold = True
    run_course.font.color.rgb = COLOR_SECONDARY

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    run_title = title_para.add_run(module_title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(12)
    run_sub = sub_para.add_run(f"Comprehensive Identification & Multiple-Choice Reviewer — {len(items)} Items")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = COLOR_GRAY

    # Student Info Table
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Inches(4.5)
    info_table.columns[1].width = Inches(2.5)

    cell_0 = info_table.cell(0, 0)
    p0 = cell_0.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    r_name = p0.add_run("Name: _____________________________________________")
    r_name.font.name = 'Arial'
    r_name.font.size = Pt(9.5)

    cell_1 = info_table.cell(0, 1)
    p1 = cell_1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r_date = p1.add_run("Date: ____________ Score: _______")
    r_date.font.name = 'Arial'
    r_date.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Instructions Box
    inst_table = doc.add_table(rows=1, cols=1)
    inst_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    inst_cell = inst_table.cell(0, 0)
    set_cell_background(inst_cell, HEX_LIGHT_BG)
    set_cell_margins(inst_cell, top=100, bottom=100, left=140, right=140)
    inst_p = inst_cell.paragraphs[0]
    inst_p.paragraph_format.space_after = Pt(0)
    r_inst_lbl = inst_p.add_run("INSTRUCTIONS: ")
    r_inst_lbl.bold = True
    r_inst_lbl.font.size = Pt(9.5)
    r_inst_lbl.font.name = 'Arial'
    r_inst_lbl.font.color.rgb = COLOR_PRIMARY

    r_inst = inst_p.add_run(
        "Read each statement carefully. Identify the correct term or concept described in the blank (_____). "
        "Choose the letter of the correct answer from the choices provided (a, b, c, or d) and write it on the blank provided. "
        "An exhaustive Answer Key with detailed rationales is provided at the end of this document."
    )
    r_inst.font.size = Pt(9)
    r_inst.font.name = 'Arial'

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section Heading
    sec_p = doc.add_paragraph()
    sec_p.paragraph_format.space_before = Pt(10)
    sec_p.paragraph_format.space_after = Pt(6)
    r_sec = sec_p.add_run(f"PART I: IDENTIFICATION QUESTIONNAIRE ({len(items)} ITEMS)")
    r_sec.font.name = 'Arial'
    r_sec.font.size = Pt(11)
    r_sec.font.bold = True
    r_sec.font.color.rgb = COLOR_PRIMARY

    # Render items
    processed_items = []
    for idx, item in enumerate(items, start=1):
        choices = [item["a"]] + item["distractors"]
        # Deterministic shuffle seeded by question index
        rng = random.Random(idx * 7919)
        rng.shuffle(choices)
        correct_letter = ['a', 'b', 'c', 'd'][choices.index(item["a"])]
        processed_items.append({
            "num": idx,
            "q": item["q"],
            "choices": choices,
            "correct_letter": correct_letter,
            "correct_text": item["a"],
            "topic": item["topic"],
            "explanation": item["explanation"]
        })

        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(4)
        q_para.paragraph_format.space_after = Pt(2)
        q_para.paragraph_format.line_spacing = 1.15

        r_num = q_para.add_run(f"_____ {idx}. ")
        r_num.bold = True
        r_num.font.name = 'Arial'
        r_num.font.size = Pt(9.5)
        r_num.font.color.rgb = COLOR_DARK

        r_text = q_para.add_run(item["q"])
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(9.5)

        ch_para = doc.add_paragraph()
        ch_para.paragraph_format.left_indent = Inches(0.4)
        ch_para.paragraph_format.space_before = Pt(1)
        ch_para.paragraph_format.space_after = Pt(4)

        letters = ['a', 'b', 'c', 'd']
        for l_idx, ch in enumerate(choices):
            r_c = ch_para.add_run(f"{letters[l_idx]}.) {ch}     ")
            r_c.font.name = 'Arial'
            r_c.font.size = Pt(9)
            r_c.font.color.rgb = COLOR_DARK

    # Answer Key Section
    doc.add_page_break()
    ak_head = doc.add_paragraph()
    ak_head.paragraph_format.space_before = Pt(12)
    ak_head.paragraph_format.space_after = Pt(4)
    r_ak = ak_head.add_run("COMPLETE ANSWER KEY & EXPLANATIONS")
    r_ak.font.name = 'Arial'
    r_ak.font.size = Pt(14)
    r_ak.font.bold = True
    r_ak.font.color.rgb = COLOR_PRIMARY

    sub_ak = doc.add_paragraph()
    sub_ak.paragraph_format.space_after = Pt(10)
    r_sub_ak = sub_ak.add_run(f"{module_title} — Reference Key with Concept Explanations")
    r_sub_ak.font.name = 'Arial'
    r_sub_ak.font.size = Pt(9.5)
    r_sub_ak.font.color.rgb = COLOR_GRAY

    # Table for Answer Key
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.6)  # Item #
    table.columns[1].width = Inches(0.6)  # Answer
    table.columns[2].width = Inches(1.8)  # Correct Term
    table.columns[3].width = Inches(4.0)  # Concept & Explanation

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Item", "Key", "Correct Answer", "Curriculum Explanation & Reference"]
    for i, t in enumerate(hdr_titles):
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(t)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for p_item in processed_items:
        row_cells = table.add_row().cells
        for cell in row_cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

        # Item #
        p0 = row_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(p_item["num"]))
        r0.font.name = 'Arial'
        r0.font.size = Pt(8.5)
        r0.font.bold = True

        # Key
        p1 = row_cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(p_item["correct_letter"].upper())
        r1.font.name = 'Arial'
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_PRIMARY

        # Correct Answer
        p2 = row_cells[2].paragraphs[0]
        r2 = p2.add_run(p_item["correct_text"])
        r2.font.name = 'Arial'
        r2.font.size = Pt(8.5)
        r2.font.bold = True

        # Explanation
        p3 = row_cells[3].paragraphs[0]
        r_top = p3.add_run(f"[{p_item['topic']}] ")
        r_top.font.name = 'Arial'
        r_top.font.size = Pt(8.5)
        r_top.font.bold = True
        r_top.font.color.rgb = COLOR_SECONDARY

        r_exp = p3.add_run(p_item["explanation"])
        r_exp.font.name = 'Arial'
        r_exp.font.size = Pt(8.5)

    set_table_borders(table)

    doc.save(output_docx_path)
    print(f"Saved DOCX: {output_docx_path}")
    return processed_items

# Run generation
out_dir = "/home/javvii/YearIII/NETC311/quiz1"
docx_file = os.path.join(out_dir, "Module 1 - Human Biology - Scientific Method and Basic Chemistry.docx")
pdf_file = os.path.join(out_dir, "Module 1 - Human Biology - Scientific Method and Basic Chemistry.pdf")

processed = build_hmby_doc(
    "Module 1: The Scientific Method & Basic Chemistry",
    "HMBY311 - Human Biology Reviewer",
    "HMBY311-M1",
    HMBY_MODULE_1_ITEMS,
    docx_file
)

# Convert to PDF via libreoffice
subprocess.run([
    'libreoffice', '--headless', '--convert-to', 'pdf',
    docx_file, '--outdir', out_dir
], check=True)

print("PDF successfully generated!")
